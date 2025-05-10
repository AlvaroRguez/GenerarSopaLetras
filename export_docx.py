# export_docx.py

import io
import math
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from tqdm import tqdm

from config import *
from drawing import draw_puzzle, draw_solution

def create_docx(all_puzzles, name: str = f"{TOTAL_PUZZLES}_word_search_puzzles.docx"):
    doc = Document()
    # cover
    para = doc.add_heading(TITLE_DOCX, level=1)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # puzzles
    for idx, (puzzle, words, _) in enumerate(
        tqdm(all_puzzles, desc="DOCX: puzzles", unit="puzzle", ncols=TQDM_COLS, position=0, leave=True),
        start=1
    ):
        para = doc.add_heading(f'{TITLE_DOCX} Nº: {idx} [{len(words)}]', level=DOCX_TITLE_LEVEL)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # image
        fig = plt.figure(figsize=(PUZZLE_COLUMNS/2, PUZZLE_ROWS/2))
        ax = fig.add_axes([0,0,1,1])
        draw_puzzle(ax, puzzle, PDF_PUZZLE_FONT)
        buf = io.BytesIO()
        fig.savefig(buf, format='png', bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)

        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture(buf, width=Inches(DOCX_IMAGE_WIDTH))
        p.alignment = DOCX_PARA_ALIGN

        # word table
        cols = SEARCH_WORDS_COLS
        rows = math.ceil(len(words)/cols)
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = DOCX_TABLE_ALIGN
        table.autofit = False
        sorted_words = sorted(words)
        for i, w in enumerate(sorted_words):
            cell = table.rows[i//cols].cells[i%cols]
            cell.text = w.upper()
            cell.paragraphs[0].alignment = DOCX_PARA_ALIGN

    # solutions
    doc.add_page_break()
    doc.add_heading('Solutions', level=1)
    doc.add_page_break()
    per_page, cols = SOLUTION_PER_PAGE, SOLUTION_COLS
    rows = math.ceil(per_page/cols)
    pages = math.ceil(len(all_puzzles)/per_page)
    for page_idx, start in enumerate(
        tqdm(range(0, len(all_puzzles), per_page),
             desc="DOCX: solutions", unit="pages",
             total=pages, ncols=TQDM_COLS),
        start=1
    ):
        if page_idx>1:
            doc.add_page_break()
        group = all_puzzles[start:start+per_page]
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = DOCX_TABLE_ALIGN
        table.autofit = False

        for i,(puz,_,locs) in enumerate(group):
            fig = plt.figure(figsize=(3,2.5))
            ax = fig.add_axes([0,0,1,1])
            draw_solution(ax, puz, locs)
            ax.text(-0.1,0.5,f"Puzzle {start+i+1}",
                    va='center',ha='right',rotation=90,
                    fontsize=PDF_WORDLIST_FONT,
                    transform=ax.transAxes)
            buf = io.BytesIO()
            fig.savefig(buf, format='png', bbox_inches='tight')
            plt.close(fig)
            buf.seek(0)

            cell = table.rows[i//cols].cells[i%cols]
            run = cell.paragraphs[0].add_run()
            run.add_picture(buf, width=Inches(DOCX_SOL_IMG_WIDTH))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_HI
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(name)
    tqdm.write(f"Word document generated: {name}")
