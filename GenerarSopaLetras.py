import random
import sys
import io
import math
import json
import time
import os

# Check all external dependencies
try:
    from tqdm import tqdm
except ImportError:
    print("Missing 'tqdm' library. Install it with: pip install tqdm", file=sys.stderr)
    sys.exit(1)

try:
    import spacy
except ImportError:
    print("Missing 'spacy' library. Install it with: pip install spacy", file=sys.stderr)
    print("You will also need to install the Spanish model: python -m spacy download es_core_news_sm", file=sys.stderr)
    sys.exit(1)

try:
    from wordfreq import top_n_list
except ImportError:
    print("Missing 'wordfreq' library. Install it with: pip install wordfreq", file=sys.stderr)
    sys.exit(1)

try:
    import matplotlib.pyplot as plt
    from matplotlib.backends.backend_pdf import PdfPages
except ImportError:
    print("Missing 'matplotlib' library. Install it with: pip install matplotlib", file=sys.stderr)
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Missing 'python-docx' library. Install it with: pip install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    import enchant  # Make sure you have pyenchant installed and the es_ES dictionary configured
except ImportError:
    print("Missing 'pyenchant' library. Install it with: pip install pyenchant", file=sys.stderr)
    print("You will also need to install the Spanish dictionary for your operating system.", file=sys.stderr)
    sys.exit(1)

# ——— GENERAL CONFIGURATION ———

# Word source: "wordfreq" or "file"
WORD_SOURCE         = "file"
WORD_SOURCE_FILE    = "palabras_todas.txt"  
MAX_RAW_WORDS       = 100000

# — Initial list filtering ——
MIN_WORD_LENGTH     = 4
MAX_WORD_LENGTH     = 10
POS_ALLOWED         = {"NOUN", "VERB", "ADV"}
BLACKLIST_FILE      = "blacklist.json"

# — Word search generation ——
WORDS_PER_PUZZLE    = 50
PUZZLE_ROWS         = 14
PUZZLE_COLUMNS      = 17
MAX_FALLBACK_TRIES  = 200
ALPHABET            = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
DIRECTIONS          = [
    (0,1),(1,0),(0,-1),(-1,0),
    (1,1),(1,-1),(-1,1),(-1,-1),
]

# — PDF Export ——
TOTAL_PUZZLES       = 365
PDF_PAGE_SIZE       = (8.27, 11.69)  # in inches (A4)
PDF_PUZZLE_AREA     = dict(left=0.1, bottom=0.30, width=0.8, height=0.55)
SEARCH_WORDS_COLS   = 5
PDF_TITLE_FONT      = dict(size=14, weight='bold')
PDF_PUZZLE_FONT     = 12
PDF_WORDLIST_FONT   = 10

# — DOCX Export ——
SOLUTION_PER_PAGE   = 15
SOLUTION_COLS       = 3
DOCX_IMAGE_WIDTH    = 6     # inches
DOCX_SOL_IMG_WIDTH  = 2     # inches
DOCX_TITLE_LEVEL    = 2
DOCX_TABLE_ALIGN    = WD_TABLE_ALIGNMENT.CENTER
DOCX_PARA_ALIGN     = WD_PARAGRAPH_ALIGNMENT.CENTER

# — Progress bar ——
TQDM_NCOLS          = 80

try:
    nlp = spacy.load("es_core_news_sm", disable=["parser","ner","lemmatizer"])
except OSError:
    print("Install the Spanish model: python -m spacy download es_core_news_sm", file=sys.stderr)
    sys.exit(1)

try:
    nlp = spacy.load("es_core_news_sm", disable=["parser", "ner", "lemmatizer"])
except OSError:
    print("Error loading spaCy model 'es_core_news_sm'. Make sure to install it with: python -m spacy download es_core_news_sm", file=sys.stderr)
    sys.exit(1)

try:
    blacklist = load_blacklist_from_json('blacklist.json')
except Exception as e:
    print(f"Warning when loading blacklist.json: {e}. An empty exclusion list will be used.", file=sys.stderr)
    blacklist = set()

def load_blacklist_from_json(filename: str) -> set[str]:
    """Loads the blacklist from a JSON file."""
    blacklist_set: set[str] = set()
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if isinstance(data, (list, set)):
                blacklist_set.update(w.lower() for w in data)
            else:
                print(f"Warning: '{filename}' does not contain a valid list or set.")
    except FileNotFoundError:
        print(f"Error: File '{filename}' not found.")
    except json.JSONDecodeError:
        print(f"Error: File '{filename}' is not a valid JSON.")
    return blacklist_set

def is_word_allowed(word: str, blacklist_input: set[str]) -> bool:
    """Checks if the word is not in the blacklist."""
    return word.lower() not in blacklist_input

def get_raw_word_list() -> list[str]:
    """Returns the raw word list according to WORD_SOURCE."""
    if WORD_SOURCE == "wordfreq":
        return top_n_list("es", MAX_RAW_WORDS)
    elif WORD_SOURCE == "file":
        try:
            with open(WORD_SOURCE_FILE, encoding="utf-8") as f:
                return [w.strip() for w in f if w.strip()]
        except FileNotFoundError:
            print(f"Error: file {WORD_SOURCE_FILE} does not exist", file=sys.stderr)
            sys.exit(1)
    else:
        print(f"WARNING: unknown source '{WORD_SOURCE}', using wordfreq.", file=sys.stderr)
        return top_n_list("es", MAX_RAW_WORDS)

def generate_word_search(
    words: list[str],
    rows: int = PUZZLE_ROWS,
    columns: int = PUZZLE_COLUMNS
) -> tuple[list[list[str]], dict[str,tuple[tuple[int,int],tuple[int,int]]]]:
    """Generates a word search puzzle using constants DIRECTIONS, ALPHABET, MAX_FALLBACK_TRIES..."""
    # 0) Sort from longest to shortest
    words = sorted(words, key=lambda w: -len(w))
    puzzle = [['' for _ in range(columns)] for _ in range(rows)]
    dir_counts = {d: 0 for d in DIRECTIONS}
    locations: dict[str, tuple[tuple[int,int], tuple[int,int]]] = {}

    for word in words:
        p = word.upper()
        candidates: list[tuple[int,int,int,int,int]] = []

        # 1) Explore valid positions
        for df, dc in DIRECTIONS:
            for r0 in range(rows):
                for c0 in range(columns):
                    rf = r0 + df*(len(p)-1)
                    cf = c0 + dc*(len(p)-1)
                    if not (0 <= rf < rows and 0 <= cf < columns):
                        continue
                    match_count = 0
                    ok = True
                    r, c = r0, c0
                    for l in p:
                        if puzzle[r][c] == l:
                            match_count += 1
                        elif puzzle[r][c] != '':
                            ok = False
                            break
                        r += df; c += dc
                    if ok:
                        candidates.append((match_count, r0, c0, df, dc))

        # 2) Choose best candidate (crossings + direction balancing)
        if candidates:
            candidates.sort(key=lambda x: x[0], reverse=True)
            max_match = candidates[0][0]
            top = [c for c in candidates if c[0] == max_match]
            top.sort(key=lambda x: dir_counts[(x[3], x[4])])
            _, r0, c0, df, dc = top[0]
            rf = r0 + df*(len(p)-1)
            cf = c0 + dc*(len(p)-1)
        else:
            # random fallback without mandatory crossing
            placed = False
            for _ in range(MAX_FALLBACK_TRIES):
                df, dc = random.choice(DIRECTIONS)
                r0, c0 = random.randrange(rows), random.randrange(columns)
                rf = r0 + df*(len(p)-1)
                cf = c0 + dc*(len(p)-1)
                if not (0 <= rf < rows and 0 <= cf < columns):
                    continue
                ok = True
                r, c = r0, c0
                for l in p:
                    if puzzle[r][c] not in ('', l):
                        ok = False
                        break
                    r += df; c += dc
                if ok:
                    placed = True
                    break
            if not placed:
                continue

        # 3) Place the word and update direction counter
        r, c = r0, c0
        for l in p:
            puzzle[r][c] = l
            r += df; c += dc
        locations[p] = ((r0, c0), (rf, cf))
        dir_counts[(df, dc)] += 1

    # 4) Fill empty spaces
    for i in range(rows):
        for j in range(columns):
            if puzzle[i][j] == '':
                puzzle[i][j] = random.choice(ALPHABET)

    return puzzle, locations

def draw_puzzle(ax, puzzle: list[list[str]]) -> None:
    """Draws the word search grid."""
    rows, cols = len(puzzle), len(puzzle[0])
    ax.axis('off')
    ax.set_xlim(0, cols)
    ax.set_ylim(0, rows)
    ax.set_aspect('equal')
    # outer border
    ax.plot([0, cols], [0, 0], color='black')
    ax.plot([cols, cols], [0, rows], color='black')
    ax.plot([cols, 0], [rows, rows], color='black')
    ax.plot([0, 0], [rows, 0], color='black')
    # letters
    for i in range(rows):
        for j in range(cols):
            ax.text(j + 0.5, rows - i - 0.5,
                    puzzle[i][j],
                    va='center', ha='center',
                    fontsize=PDF_PUZZLE_FONT)

def draw_solution(ax, puzzle: list[list[str]],
                  locations: dict[str,tuple[tuple[int,int],tuple[int,int]]]
) -> None:
    """Draws the solution with red lines."""
    rows, cols = len(puzzle), len(puzzle[0])
    draw_puzzle(ax, puzzle)  # reuse the grid
    for word, ((r0,c0),(rf,cf)) in locations.items():
        x0, y0 = c0 + 0.5, rows - r0 - 0.5
        x1, y1 = cf + 0.5, rows - rf - 0.5
        ax.plot([x0, x1], [y0, y1], color='red', linewidth=2)

def create_pdf(all_puzzles, name: str = f"{TOTAL_PUZZLES}_word_search_puzzles.pdf"):
    with PdfPages(name) as pp:
        for idx, (puzzle, words, _) in enumerate(
            tqdm(all_puzzles, desc="PDF: puzzles", unit="puzzle", ncols=TQDM_NCOLS),
            start=1
        ):
            fig = plt.figure(figsize=PDF_PAGE_SIZE)
            # title
            fig.text(
                0.5, 1 - PDF_PUZZLE_AREA['bottom'] + 0.02,
                f"Puzzle {idx} – {len(words)} words",
                ha='center', va='top',
                fontsize=PDF_TITLE_FONT['size'],
                weight=PDF_TITLE_FONT['weight']
            )
            # grid
            ax = fig.add_axes([
                PDF_PUZZLE_AREA['left'],
                PDF_PUZZLE_AREA['bottom'],
                PDF_PUZZLE_AREA['width'],
                PDF_PUZZLE_AREA['height']
            ])
            draw_puzzle(ax, puzzle)
            # word list
            cols, per_col = SEARCH_WORDS_COLS, math.ceil(len(words)/SEARCH_WORDS_COLS)
            gap = PDF_PUZZLE_AREA['width'] / cols
            x_positions = [PDF_PUZZLE_AREA['left'] + gap/2 + i*gap for i in range(cols)]
            y_start, y_step = PDF_PUZZLE_AREA['bottom'] - 0.02, PDF_PUZZLE_AREA['height']/per_col
            for i, w in enumerate(words):
                col = i % cols
                row = i // cols
                fig.text(
                    x_positions[col],
                    y_start - row*y_step,
                    w.upper(),
                    va='top', ha='center',
                    fontsize=PDF_WORDLIST_FONT
                )
            pp.savefig(fig)
            plt.close(fig)

        # solution pages
        per_page = SOLUTION_PER_PAGE
        pages = math.ceil(len(all_puzzles)/per_page)
        for page_idx, start in enumerate(
            tqdm(range(0, len(all_puzzles), per_page),
                 desc="PDF: solutions", unit="pages",
                 total=pages, ncols=TQDM_NCOLS),
            start=1
        ):
            fig = plt.figure(figsize=PDF_PAGE_SIZE)
            for sub, loc in enumerate(range(start, min(start+per_page, len(all_puzzles)))):
                puzzle, _, locs = all_puzzles[loc]
                col = sub % 2
                row = sub // 2
                left = 0.05 + col*0.475
                bottom = 0.05 + (4-row)*0.18
                ax = fig.add_axes([left, bottom, 0.45, 0.18])
                draw_solution(ax, puzzle, locs)
                ax.text(-0.1, 0.5, f"Puzzle {loc+1}",
                        va='center', ha='right', rotation=90,
                        fontsize=PDF_WORDLIST_FONT,
                        transform=ax.transAxes)
            pp.savefig(fig)
            plt.close(fig)

    print(f"PDF generated: {name}")

def create_docx(all_puzzles, name: str = f"{TOTAL_PUZZLES}_word_search_puzzles.docx"):
    doc = Document()
    # cover
    doc.add_heading('Word Search Puzzles', level=1)
    doc.add_page_break()

    # puzzles
    for idx, (puzzle, words, _) in enumerate(
        tqdm(all_puzzles, desc="DOCX: puzzles", unit="puzzle", ncols=TQDM_NCOLS),
        start=1
    ):
        doc.add_heading(f'Puzzle {idx} – {len(words)} words', level=DOCX_TITLE_LEVEL)
        # image
        fig = plt.figure(figsize=(PUZZLE_COLUMNS/2, PUZZLE_ROWS/2))
        ax = fig.add_axes([0,0,1,1])
        draw_puzzle(ax, puzzle)
        buf = io.BytesIO()
        fig.savefig(buf, format='png', bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)

        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture(buf, width=Inches(DOCX_IMAGE_WIDTH))
        p.alignment = DOCX_PARA_ALIGN

        # word table
        cols = SEARCH_WORDS_COLS
        rows = math.ceil(len(words)/cols)
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = DOCX_TABLE_ALIGN
        table.autofit = False
        for i, w in enumerate(words):
            cell = table.rows[i//cols].cells[i%cols]
            cell.text = w.upper()
            cell.paragraphs[0].alignment = DOCX_PARA_ALIGN

    # solutions
    doc.add_page_break()
    doc.add_heading('Solutions', level=1)
    per_page, cols = SOLUTION_PER_PAGE, SOLUTION_COLS
    rows = math.ceil(per_page/cols)
    pages = math.ceil(len(all_puzzles)/per_page)
    for page_idx, start in enumerate(
        tqdm(range(0, len(all_puzzles), per_page),
             desc="DOCX: solutions", unit="pages",
             total=pages, ncols=TQDM_NCOLS),
        start=1
    ):
        if page_idx>1:
            doc.add_page_break()
        group = all_puzzles[start:start+per_page]
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = DOCX_TABLE_ALIGN
        table.autofit = False

        for i,(puz,_,locs) in enumerate(group):
            fig = plt.figure(figsize=(3,2.5))
            ax = fig.add_axes([0,0,1,1])
            draw_solution(ax, puz, locs)
            ax.text(-0.1,0.5,f"Puzzle {start+i+1}",
                    va='center',ha='right',rotation=90,
                    fontsize=PDF_WORDLIST_FONT,
                    transform=ax.transAxes)
            buf = io.BytesIO()
            fig.savefig(buf, format='png', bbox_inches='tight')
            plt.close(fig)
            buf.seek(0)

            cell = table.rows[i//cols].cells[i%cols]
            run = cell.paragraphs[0].add_run()
            run.add_picture(buf, width=Inches(DOCX_SOL_IMG_WIDTH))
            cell.paragraphs[0].alignment = DOCX_PARA_ALIGN

    doc.save(name)
    print(f"Word document generated: {name}")

def build_filtered_dict(raw: list[str], blacklist_input: set[str]) -> list[str]:
    """Applies prefiltering by length, characters and POS with spaCy."""
    pre = [
        w for w in raw
        if MIN_WORD_LENGTH <= len(w) <= MAX_WORD_LENGTH
           and w.isascii()
           and w.isalpha()
           and w.lower() not in blacklist_input
    ]
    print(f"  → After prefiltering: {len(pre)} words for POS.")

    filtered: list[str] = []
    for doc in nlp.pipe(pre, batch_size=2000, n_process=1):
        tok = doc[0]
        if tok.pos_ in POS_ALLOWED:
            filtered.append(tok.text)
    return filtered

def main():
    raw = get_raw_word_list()
    print(f"Retrieved {len(raw)} raw words.")

    blacklist = load_blacklist_from_json(BLACKLIST_FILE)
    filtered_dict = build_filtered_dict(raw, blacklist)
    print(f"Filtered dictionary: {len(filtered_dict)} valid words.\n")

    # save filtered dictionary
    out = f"{WORD_SOURCE}_{WORD_SOURCE_FILE}_filtered.txt".replace(os.sep, "_")
    with open(out, "w", encoding="utf-8") as f:
        for w in filtered_dict:
            f.write(w + "\n")
    print(f"Filtered dictionary saved in '{out}'\n")

    # generate puzzles
    all_puzzles = []
    used: set[str] = set()
    start = time.perf_counter()
    pbar = tqdm(range(TOTAL_PUZZLES), desc="Generating puzzles", unit="puzzle", ncols=TQDM_NCOLS)
    for _ in pbar:
        avail = [w for w in filtered_dict if w.upper() not in used]
        if len(avail) < WORDS_PER_PUZZLE:
            used.clear()
            avail = filtered_dict.copy()
        sel = random.sample(avail, WORDS_PER_PUZZLE)
        used.update(w.upper() for w in sel)
        puz, locs = generate_word_search(sel)
        all_puzzles.append((puz, [w.upper() for w in sel], locs))
        pbar.set_postfix({"elapsed": f"{time.perf_counter()-start:.1f}s"})
    pbar.close()

    print("\nGeneration completed.\nCreating PDF and DOCX...")
    create_docx(all_puzzles)
    create_pdf(all_puzzles)
    print("Process finished!")

if __name__ == '__main__':
    main()
